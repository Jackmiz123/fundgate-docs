#!/usr/bin/env python3
"""Generate FundGate Payoff Letter as DOCX bytes."""
import zipfile, io, re, os
from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

LOGO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                         'FUNDGATE_TEMPLATE_WEEKLY.docx')


def _extract_logo_bytes():
    """Pull the FundGate logo image from the weekly template."""
    with zipfile.ZipFile(LOGO_PATH) as z:
        # image1.jpeg is the FundGate logo
        return z.read('word/media/image1.jpeg')


def build_payoff_letter(data):
    """Build a payoff letter DOCX and return bytes.

    Required data keys:
        payoff_date        – e.g. '04/16/2026'
        payoff_merchant    – e.g. 'Beck Inc. Pipeline Construction'
        payoff_balance     – e.g. '$33,900.00'
        payoff_wire_amount – e.g. '$28,650.00'
    """
    date_str = data.get('payoff_date', '')
    merchant = data.get('payoff_merchant', '')
    original_amt = data.get('payoff_original_amount', '')  # optional
    balance = data.get('payoff_balance', '')
    wire_amt = data.get('payoff_wire_amount', '')

    if not all([date_str, merchant, balance, wire_amt]):
        return None

    # Format date nicely: 04/16/2026 -> April 16, 2026
    display_date = date_str
    try:
        parts = date_str.split('/')
        if len(parts) == 3:
            from datetime import datetime
            dt = datetime(int(parts[2]), int(parts[0]), int(parts[1]))
            display_date = dt.strftime('%B %d, %Y')
    except Exception:
        pass

    doc = Document()

    # ── Page margins ──
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # ── Logo ──
    logo_bytes = _extract_logo_bytes()
    logo_stream = io.BytesIO(logo_bytes)
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_para.space_after = Pt(20)
    logo_run = logo_para.add_run()
    logo_run.add_picture(logo_stream, width=Inches(3.2))

    # ── Horizontal line ──
    line_para = doc.add_paragraph()
    line_para.space_before = Pt(0)
    line_para.space_after = Pt(16)
    pPr = line_para._element.get_or_add_pPr()
    from lxml import etree
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    pBdr = etree.SubElement(pPr, f'{{{nsmap["w"]}}}pBdr')
    bottom = etree.SubElement(pBdr, f'{{{nsmap["w"]}}}bottom')
    bottom.set(f'{{{nsmap["w"]}}}val', 'single')
    bottom.set(f'{{{nsmap["w"]}}}sz', '6')
    bottom.set(f'{{{nsmap["w"]}}}space', '1')
    bottom.set(f'{{{nsmap["w"]}}}color', '000000')

    # ── Date ──
    date_para = doc.add_paragraph(display_date)
    date_para.space_after = Pt(8)
    for run in date_para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    # ── Merchant Name (bold) ──
    merch_para = doc.add_paragraph()
    merch_para.space_after = Pt(4)
    merch_run = merch_para.add_run(merchant)
    merch_run.bold = True
    merch_run.font.name = 'Times New Roman'
    merch_run.font.size = Pt(12)

    # ── "PAYOFF LETTER" heading ──
    heading_para = doc.add_paragraph()
    heading_para.space_after = Pt(16)
    h_run = heading_para.add_run('PAYOFF LETTER')
    h_run.bold = True
    h_run.underline = True
    h_run.font.name = 'Times New Roman'
    h_run.font.size = Pt(12)

    # ── Body paragraph 1 ──
    def _add_body_para(doc_obj):
        p = doc_obj.add_paragraph()
        p.paragraph_format.space_after = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    def _add_run(para, text, bold=False):
        r = para.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.bold = bold
        return r

    p1 = _add_body_para(doc)
    if original_amt:
        _add_run(p1, 'The original purchased amount on your account is ')
        _add_run(p1, original_amt, bold=True)
        _add_run(p1, '. ')
    _add_run(p1, 'The balance of your account as of the date above is ')
    _add_run(p1, balance, bold=True)
    _add_run(p1, '. Please note this balance is subject to change in the event of an ACH debit '
                 'payment being rejected for any reason. If there is a rejection of a payment, you '
                 'will be subject to the appropriate fees detailed in Schedule A of the Merchant '
                 'Agreement in addition to the payment that was rejected.')

    # ── Body paragraph 2 ──
    p2 = _add_body_para(doc)
    _add_run(p2, 'If you choose to pay this balance today, please wire ')
    _add_run(p2, wire_amt, bold=True)
    _add_run(p2, ' to the following account:')

    # ── Bank details (indented, bold) ──
    bank_lines = [
        'Optimum Bank',
        'FundGate LLC',
        '2929 E Commercial Boulevard',
        'Fort Lauderdale, FL 33308',
        'ABA # 067015096',
        'Acct# 210058947',
    ]
    for line in bank_lines:
        bp = doc.add_paragraph()
        bp.paragraph_format.left_indent = Inches(0.75)
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after = Pt(0)
        br = bp.add_run(line)
        br.bold = True
        br.font.name = 'Times New Roman'
        br.font.size = Pt(12)
    # add spacing after bank block
    doc.paragraphs[-1].paragraph_format.space_after = Pt(16)

    # ── PLEASE NOTE disclaimer ──
    note_para = doc.add_paragraph()
    note_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    note_para.paragraph_format.space_after = Pt(16)
    note_run = note_para.add_run(
        'PLEASE NOTE: IF PAID OFF BY A THIRD PARTY, THIS DISCOUNT WILL BE '
        'CONSIDERED INVALID.'
    )
    note_run.bold = True
    note_run.font.name = 'Times New Roman'
    note_run.font.size = Pt(12)

    # ── Balance adjustments paragraph ──
    p3 = _add_body_para(doc)
    _add_run(p3, 'Any balance adjustments or amounts which are currently being processed and collected '
                 'by FundGate LLC after the remitted payoff balance has cleared our account will be '
                 'immediately returned to your designated operating account.')

    # ── Closing ──
    close1 = doc.add_paragraph()
    close1.paragraph_format.space_before = Pt(24)
    close1.paragraph_format.space_after = Pt(12)
    _add_run(close1, 'Sincerely,')

    for line in ['Accounts Receivable', 'FundGate LLC', 'admin@fundgatellc.com']:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(0)
        _add_run(cp, line)

    # ── Save to bytes ──
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
