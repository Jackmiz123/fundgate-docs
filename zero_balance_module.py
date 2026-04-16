#!/usr/bin/env python3
"""Generate FundGate Zero Balance Letter as DOCX bytes."""
import zipfile, io, os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

LOGO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                         'FUNDGATE_TEMPLATE_WEEKLY.docx')


def _extract_logo_bytes():
    with zipfile.ZipFile(LOGO_PATH) as z:
        return z.read('word/media/image1.jpeg')


def build_zero_balance_letter(data):
    """Build a zero balance letter DOCX and return bytes.

    Required data keys:
        zb_date     – e.g. '04/15/2026'
        zb_merchant – e.g. 'J&J Roofing LLC'
    """
    date_str = data.get('zb_date', '')
    merchant = data.get('zb_merchant', '')

    if not all([date_str, merchant]):
        return None

    # Format date: 04/15/2026 -> April 15, 2026
    display_date = date_str
    try:
        parts = date_str.split('/')
        if len(parts) == 3:
            from datetime import datetime
            dt = datetime(int(parts[2]), int(parts[0]), int(parts[1]))
            display_date = dt.strftime('%B %d, %Y')
    except Exception:
        pass

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    def _add_run(para, text, bold=False, underline=False):
        r = para.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.bold = bold
        r.underline = underline
        return r

    # ── Logo ──
    logo_bytes = _extract_logo_bytes()
    logo_stream = io.BytesIO(logo_bytes)
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_para.space_after = Pt(20)
    logo_run = logo_para.add_run()
    logo_run.add_picture(logo_stream, width=Inches(3.2))

    # ── Horizontal line ──
    line_para = doc.add_paragraph()
    line_para.space_before = Pt(0)
    line_para.space_after = Pt(16)
    pPr = line_para._element.get_or_add_pPr()
    from lxml import etree
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    pBdr = etree.SubElement(pPr, f'{{{nsmap["w"]}}}pBdr')
    bottom = etree.SubElement(pBdr, f'{{{nsmap["w"]}}}bottom')
    bottom.set(f'{{{nsmap["w"]}}}val', 'single')
    bottom.set(f'{{{nsmap["w"]}}}sz', '6')
    bottom.set(f'{{{nsmap["w"]}}}space', '1')
    bottom.set(f'{{{nsmap["w"]}}}color', '000000')

    # ── Date ──
    date_para = doc.add_paragraph()
    date_para.space_after = Pt(8)
    _add_run(date_para, display_date)

    # ── Merchant Name (bold) ──
    merch_para = doc.add_paragraph()
    merch_para.space_after = Pt(4)
    _add_run(merch_para, merchant, bold=True)

    # ── "ZERO BALANCE LETTER" heading ──
    heading_para = doc.add_paragraph()
    heading_para.space_after = Pt(16)
    _add_run(heading_para, 'ZERO BALANCE LETTER', bold=True, underline=True)

    # ── Body paragraph ──
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1.paragraph_format.space_after = Pt(16)
    _add_run(p1, 'The balance of your account as of the date above is ')
    _add_run(p1, '$0.00', bold=True)
    _add_run(p1, '. Please note this balance is subject to change in the event of an ACH debit '
                 'payment being rejected for any reason. If there is a rejection of a payment, you '
                 'will be subject to the appropriate fees detailed in Schedule A of the Merchant '
                 'Agreement in addition to the payment that was rejected.')

    # ── Closing ──
    close1 = doc.add_paragraph()
    close1.paragraph_format.space_before = Pt(24)
    close1.paragraph_format.space_after = Pt(12)
    _add_run(close1, 'Sincerely,')

    for line in ['Accounts Receivable', 'FundGate LLC', 'admin@fundgatellc.com', '929-256-7464']:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(0)
        _add_run(cp, line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
